from typing import Any, List
import docx, re, difflib
from collections import namedtuple


def __is_subitem_of_contract(text: str) -> bool:
    return re.match("(\d+\.){2,5}", text.strip())

def __is_item_of_contract(text: str) -> bool:
    return re.match("(\d+\.){1}", text.strip())

def __is_continious(text: str) -> bool:
    return text.strip()

def __add_specified_run(parag: Any, text: str, *, bold: bool = False, underline: bool = False, italic: bool = False, colorRGB: docx.shared.RGBColor = None, strike: bool = None) -> None:
    run = parag.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = colorRGB
    run.font.strike = strike

def __get_contract_text(parags: Any, endmark: str) -> str:
    txt = ""
    for p in parags:
        if __is_subitem_of_contract(p.text):
            txt += f"\t{p.text}\n"
        elif __is_item_of_contract(p.text):
            if p.text == endmark:
                break
            else: txt += f"\n{p.text}\n\n"
        elif __is_continious(p.text):
            txt += f"{p.text}\n"
    return txt

def make_report_docx(fragments: List[tuple], path_to_save: str) -> None:

    doc_report = docx.Document()
    parag = doc_report.add_paragraph("") # All in one paragraph
    
    current_style = parag.style
    current_style.font.name = "Arial"
    current_style.font.size = docx.shared.Pt(10)
    
    INSERT_FONT = docx.shared.RGBColor(0x42, 0x24, 0xE9)
    REPLACE_FONT = docx.shared.RGBColor(0xFF, 0x88, 0x00)
    DELETE_FONT = docx.shared.RGBColor(0xFF, 0x00, 0x00)

    for fragment in fragments:
        if fragment.tag == "equal":
            __add_specified_run(parag, fragment.old_text)
        elif fragment.tag == "insert":
            __add_specified_run(parag, f"{fragment.new_text}", bold=True, colorRGB=INSERT_FONT)
        elif fragment.tag == "delete":
            __add_specified_run(parag, f"{fragment.old_text}", bold=True, colorRGB=DELETE_FONT, strike=True)
        elif fragment.tag == "replace":
            __add_specified_run(parag, f"{fragment.old_text}", bold=True, colorRGB=REPLACE_FONT, strike=True)
            __add_specified_run(parag, f" ---> ", bold=True, colorRGB=REPLACE_FONT)
            __add_specified_run(parag, f"{fragment.new_text}", bold=True, colorRGB=REPLACE_FONT)

    doc_report.save(path_to_save)

def get_marked_fragments(pattern_path: str, edited_path: str, numsym_before: int, numsym_after: int, end_mark: str) -> List[tuple]:

    pattern_txt = __get_contract_text(docx.Document(pattern_path).paragraphs, end_mark)
    edited_txt = __get_contract_text(docx.Document(edited_path).paragraphs, end_mark)
    
    seq_mtcher = difflib.SequenceMatcher(a = pattern_txt, b = edited_txt)
    fragments: List[Fragment] = [] # type: ignore
    Fragment = namedtuple("Fragment", ["tag", "txt_before", "old_text", "new_text", "txt_after"])

    for tag, i1, i2, j1, j2 in seq_mtcher.get_opcodes():

        if tag == "equal":
            fragments.append(Fragment(tag, 
                                      f"...{pattern_txt[max(0, (i1-numsym_before)):i1]}", 
                                      pattern_txt[i1:i2], 
                                      "", 
                                      f"{pattern_txt[i2:min((j2+numsym_after), len(edited_txt))]}..."))
        elif tag == "insert":
            fragments.append(Fragment(tag, 
                                      f"...{edited_txt[max(0, (j1-numsym_before)):j1]}", 
                                      "", 
                                      edited_txt[j1:j2], 
                                      f"{edited_txt[j2:min((j2+numsym_after), len(edited_txt))]}..."))
        elif tag == "delete":
            fragments.append(Fragment(tag, 
                                      f"...{pattern_txt[max(0, (i1-numsym_before)):i1]}", 
                                      pattern_txt[i1:i2], 
                                      "", 
                                      f"{pattern_txt[i2:min((i2+numsym_after), len(pattern_txt))]}..."))
        elif tag == "replace":
            fragments.append(Fragment(tag, 
                                      f"...{pattern_txt[max(0, (i1-numsym_before)):i1]}", 
                                      pattern_txt[i1:i2], 
                                      edited_txt[j1:j2], 
                                      f"{edited_txt[j2:min((j2+numsym_after), len(edited_txt))]}..."))

    return fragments
